from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login
from django.contrib.auth.backends import BaseBackend
from django.http import HttpResponseRedirect, HttpResponse
from django.db import models  # Añade esta línea
from .models import Estudiante, Docente, Nodocente, Asignatura, Evento, Asistencia, Anotacion, Nota
from .forms import LoginDocenteForm, LoginNodocenteForm, CreateEstudianteForm, CreateDocenteForm, CreateNodocenteForm, CreateAsignaturaForm, AssignAsignaturaForm, CreateEventoForm, CreateAnotacionForm, CreateNotaForm
from django.utils import timezone
from datetime import date
from django.contrib.auth.decorators import login_required
import csv
from django.http import HttpResponse
from docx import Document
import os
from django.conf import settings
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
# Resto del código...
class EstudianteBackend(BaseBackend):
    def authenticate(self, request, rut=None):
        try:
            estudiante = Estudiante.objects.get(rut=rut)
            return estudiante
        except Estudiante.DoesNotExist:
            return None

    def get_user(self, user_id):
        try:
            return Estudiante.objects.get(pk=user_id)
        except Estudiante.DoesNotExist:
            return None

class DocenteBackend(BaseBackend):
    def authenticate(self, request, correo=None, contrasena=None):
        try:
            docente = Docente.objects.get(correo=correo)
            if docente.contrasena == contrasena:
                return docente
        except Docente.DoesNotExist:
            return None

    def get_user(self, user_id):
        try:
            return Docente.objects.get(pk=user_id)
        except Docente.DoesNotExist:
            return None

class NodocenteBackend(BaseBackend):
    def authenticate(self, request, correo=None, contrasena=None):
        try:
            nodocente = Nodocente.objects.get(correo=correo)
            if nodocente.contrasena == contrasena:
                return nodocente
        except Nodocente.DoesNotExist:
            return None

    def get_user(self, user_id):
        try:
            return Nodocente.objects.get(pk=user_id)
        except Nodocente.DoesNotExist:
            return None

def home_view(request):
    return render(request, 'home.html')

def login_estudiante_view(request):
    error_message = None
    if request.method == 'POST':
        rut = request.POST['rut']
        estudiante = EstudianteBackend().authenticate(request, rut=rut)
        if estudiante:
            login(request, estudiante, backend='sistema_Colegio.backends.EstudianteBackend')
            return redirect('estudiante_menu')
        else:
            error_message = "RUT incorrecto. Por favor, inténtelo de nuevo."
    return render(request, 'estudiante_login.html', {'error_message': error_message})

def login_docente_view(request):
    error_message = None
    if request.method == 'POST':
        form = LoginDocenteForm(request.POST)
        if form.is_valid():
            correo = form.cleaned_data['correo']
            password = form.cleaned_data['password']
            if correo == 'admin@admin.com' and password == 'admin':
                # Crear o obtener el usuario admin
                docente, created = Docente.objects.get_or_create(
                    correo=correo,
                    defaults={'nombre': 'Admin', 'rut': '12345678-9', 'tipo': 'Admin', 'is_admin': True}
                )
                if created:
                    docente.set_password(password)
                    docente.save()
                login(request, docente, backend='django.contrib.auth.backends.ModelBackend')
                return redirect('admin_menu')
            else:
                docente = Docente.objects.filter(correo=correo).first()
                if docente and docente.check_password(password):
                    login(request, docente, backend='django.contrib.auth.backends.ModelBackend')
                    return redirect('docente_menu')
                else:
                    error_message = "Correo o contraseña incorrectos. Por favor, inténtelo de nuevo."
    else:
        form = LoginDocenteForm()
    return render(request, 'docente_login.html', {'form': form, 'error_message': error_message})

def login_nodocente_view(request):
    error_message = None
    if request.method == 'POST':
        form = LoginNodocenteForm(request.POST)
        if form.is_valid():
            correo = form.cleaned_data['correo']
            password = form.cleaned_data['password']
            nodocente = NodocenteBackend().authenticate(request, correo=correo, password=password)
            if nodocente:
                login(request, nodocente, backend='sistema_Colegio.backends.NodocenteBackend')
                return redirect('nodocente_menu')
            else:
                error_message = "Correo o contraseña incorrectos. Por favor, inténtelo de nuevo."
    else:
        form = LoginNodocenteForm()
    return render(request, 'nodocente_login.html', {'form': form, 'error_message': error_message})

def admin_menu_view(request):
    return render(request, 'admin_menu.html')

def create_estudiante_view(request):
    if request.method == 'POST':
        form = CreateEstudianteForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('admin_menu')
    else:
        form = CreateEstudianteForm()
    return render(request, 'create_estudiante.html', {'form': form})

def create_docente_view(request):
    if request.method == 'POST':
        form = CreateDocenteForm(request.POST)
        if form.is_valid():
            docente = form.save(commit=False)
            docente.set_password(form.cleaned_data['contrasena'])
            docente.save()
            return redirect('admin_menu')
    else:
        form = CreateDocenteForm()
    return render(request, 'create_docente.html', {'form': form})

def create_nodocente_view(request):
    if request.method == 'POST':
        form = CreateNodocenteForm(request.POST)
        if form.is_valid():
            nodocente = form.save(commit=False)
            nodocente.set_password(form.cleaned_data['password'])
            nodocente.save()
            return redirect('admin_menu')
    else:
        form = CreateNodocenteForm()
    return render(request, 'create_nodocente.html', {'form': form})

def create_asignatura_view(request):
    if request.method == 'POST':
        form = CreateAsignaturaForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('admin_menu')
    else:
        form = CreateAsignaturaForm()
    return render(request, 'create_asignatura.html', {'form': form})

def assign_asignatura_view(request):
    estudiante_id = request.GET.get('estudiante_id')
    if estudiante_id:
        estudiante = Estudiante.objects.get(id=estudiante_id)
        if request.method == 'POST':
            form = AssignAsignaturaForm(request.POST)
            if form.is_valid():
                asignaturas = form.cleaned_data['asignatura']
                for asignatura in asignaturas:
                    estudiante.asignaturas.add(asignatura)
                return redirect('admin_menu')
        else:
            form = AssignAsignaturaForm(initial={'estudiante': estudiante})
        return render(request, 'assign_asignatura.html', {'form': form, 'estudiante': estudiante})
    else:
        return redirect('select_estudiante')

def assign_docente_view(request, asignatura_id):
    asignatura = get_object_or_404(Asignatura, id=asignatura_id)
    if request.method == 'POST':
        docente_id = request.POST.get('docente')
        docente = get_object_or_404(Docente, id=docente_id)
        asignatura.docente = docente
        asignatura.save()
        return redirect('list_asignaturas')
    docentes = Docente.objects.all()
    return render(request, 'assign_docente.html', {'asignatura': asignatura, 'docentes': docentes})

@login_required
def list_estudiantes_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'list_estudiantes.html', {'estudiantes': estudiantes})

def list_asistencia_view(request):
    fecha = request.GET.get('fecha')
    asistencia = Asistencia.objects.all()

    if fecha:
        asistencia = asistencia.filter(fecha=fecha)

    return render(request, 'list_asistencia.html', {'asistencia': asistencia, 'fecha': fecha})

def list_docentes_view(request):
    docentes = Docente.objects.all()
    return render(request, 'list_docentes.html', {'docentes': docentes})

def list_eventos_view(request):
    eventos = Evento.objects.all()
    return render(request, 'list_eventos.html', {'eventos': eventos})

def select_estudiante_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante.html', {'estudiantes': estudiantes})

def create_evento_view(request):
    if request.method == 'POST':
        form = CreateEventoForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('admin_menu')
    else:
        form = CreateEventoForm()
    return render(request, 'create_evento.html', {'form': form})

@login_required
def estudiante_menu(request):
    estudiante = request.user
    return render(request, 'estudiante_menu.html', {'estudiante': estudiante})

def docente_menu(request):
    docente = request.user
    asignaturas = Asignatura.objects.filter(docente=docente)
    return render(request, 'docente_menu.html', {'asignaturas': asignaturas})

@login_required
def nodocente_menu(request):
    eventos = Evento.objects.all()
    return render(request, 'nodocente_menu.html', {'eventos': eventos})

def registro_asistencia(request, asignatura_id):
    asignatura = get_object_or_404(Asignatura, id=asignatura_id)
    estudiantes = asignatura.estudiantes.all()
    error_message = None

    if request.method == 'POST':
        fecha = request.POST.get('fecha')
        if not fecha:
            error_message = "Por favor, ingrese una fecha antes de registrar la asistencia."
        else:
            asistencia_completa = 'asistencia_completa' in request.POST
            for estudiante in estudiantes:
                presente = asistencia_completa or request.POST.get(f'estudiante_{estudiante.id}', 'off') == 'on'
                Asistencia.objects.create(estudiante=estudiante, asignatura=asignatura, fecha=fecha, presente=presente)
            return HttpResponseRedirect('/docente_menu/')

    return render(request, 'registro_asistencia.html', {'asignatura': asignatura, 'estudiantes': estudiantes, 'fecha': date.today(), 'error_message': error_message})

def select_asignatura_view(request):
    asignaturas = Asignatura.objects.all()
    return render(request, 'select_asignatura.html', {'asignaturas': asignaturas})

def asistencia_matriz_view(request, asignatura_id):
    asignatura = get_object_or_404(Asignatura, id=asignatura_id)
    estudiantes = asignatura.estudiantes.all()
    fechas = Asistencia.objects.filter(asignatura=asignatura).values_list('fecha', flat=True).distinct()
    asistencia = {fecha: {estudiante.id: False for estudiante in estudiantes} for fecha in fechas}

    for registro in Asistencia.objects.filter(asignatura=asignatura):
        asistencia[registro.fecha][registro.estudiante.id] = registro.presente

    if request.method == 'POST':
        fecha = request.POST.get('fecha')
        if fecha:
            for estudiante in estudiantes:
                presente = request.POST.get(f'estudiante_{estudiante.id}', 'off') == 'on'
                Asistencia.objects.update_or_create(
                    estudiante=estudiante,
                    asignatura=asignatura,
                    fecha=fecha,
                    defaults={'presente': presente}
                )
            return redirect('asistencia_matriz', asignatura_id=asignatura.id)

    return render(request, 'asistencia_matriz.html', {
        'asignatura': asignatura,
        'estudiantes': estudiantes,
        'fechas': sorted(fechas),
        'asistencia': asistencia,
        'fecha_hoy': date.today()
    })

@login_required
def list_asignaturas_view(request):
    asignaturas = Asignatura.objects.all()
    return render(request, 'list_asignaturas.html', {'asignaturas': asignaturas})

def select_estudiante_anotacion_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_anotacion.html', {'estudiantes': estudiantes})

def create_anotacion_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    docente = request.user
    if request.method == 'POST':
        form = CreateAnotacionForm(request.POST)
        if form.is_valid():
            anotacion = form.save(commit=False)
            anotacion.docente = docente
            anotacion.estudiante = estudiante
            anotacion.save()
            return redirect('docente_menu')
    else:
        form = CreateAnotacionForm()
    return render(request, 'create_anotacion.html', {'form': form, 'estudiante': estudiante})

def select_asignatura_nota_view(request):
    asignaturas = Asignatura.objects.all()
    return render(request, 'select_asignatura_nota.html', {'asignaturas': asignaturas})

def create_nota_view(request, asignatura_id):
    asignatura = get_object_or_404(Asignatura, id=asignatura_id)
    estudiantes = asignatura.estudiantes.all()

    if request.method == 'POST':
        for estudiante in estudiantes:
            # Prefijo único para cada estudiante
            form = CreateNotaForm(request.POST, prefix=str(estudiante.id))
            if form.is_valid():
                nota = form.save(commit=False)
                nota.estudiante = estudiante
                nota.asignatura = asignatura
                nota.save()
        return redirect('admin_menu')
    else:
        # Crear un formulario por cada estudiante
        forms = [CreateNotaForm(prefix=str(estudiante.id)) for estudiante in estudiantes]
    
    # Combinar estudiantes y formularios para enviarlos a la plantilla
    estudiantes_forms = zip(estudiantes, forms)
    return render(request, 'create_nota.html', {
        'asignatura': asignatura,
        'estudiantes_forms': estudiantes_forms,
    })

@login_required
def ver_notas_estudiante_view(request):
    estudiante = request.user
    notas = Nota.objects.filter(estudiante=estudiante)
    return render(request, 'ver_notas_estudiante.html', {'notas': notas})

def reporte_asistencia_view(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="reporte_asistencia.csv"'

    writer = csv.writer(response)
    writer.writerow(['Estudiante', 'Asignatura', 'Fecha', 'Presente'])

    asistencia = Asistencia.objects.all()
    for registro in asistencia:
        writer.writerow([registro.estudiante.nombre, registro.asignatura.nombre, registro.fecha, 'Sí' if registro.presente else 'No'])

    return response

def reporte_notas_view(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="reporte_notas.csv"'

    writer = csv.writer(response)
    writer.writerow(['Estudiante', 'Asignatura', 'Nota'])

    notas = Nota.objects.all()
    for registro in notas:
        writer.writerow([registro.estudiante.nombre, registro.asignatura.nombre, registro.nota])

    return response

def reporte_anotaciones_view(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="reporte_anotaciones.csv"'

    writer = csv.writer(response)
    writer.writerow(['Estudiante', 'Docente', 'Fecha', 'Descripción'])

    anotaciones = Anotacion.objects.all()
    for registro in anotaciones:
        writer.writerow([registro.estudiante.nombre, registro.docente.nombre, registro.fecha, registro.descripcion])

    return response

def estadisticas_rendimiento_view(request):
    estudiantes = Estudiante.objects.all()
    estadisticas = []

    for estudiante in estudiantes:
        notas = Nota.objects.filter(estudiante=estudiante)
        promedio = notas.aggregate(models.Avg('nota'))['nota__avg']
        estadisticas.append({
            'estudiante': estudiante,
            'promedio': promedio,
        })

    return render(request, 'estadisticas_rendimiento.html', {'estadisticas': estadisticas})

def list_notas_view(request, asignatura_id):
    asignatura = get_object_or_404(Asignatura, id=asignatura_id)
    notas = Nota.objects.filter(asignatura=asignatura)
    return render(request, 'list_notas.html', {'asignatura': asignatura, 'notas': notas})
    
def perfil_estudiante_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    notas = Nota.objects.filter(estudiante=estudiante).order_by('asignatura__nombre')
    return render(request, 'perfil_estudiante.html', {'estudiante': estudiante, 'notas': notas})

def select_estudiante_reporte_asistencia_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_reporte_asistencia.html', {'estudiantes': estudiantes})

def reporte_asistencia_estudiante_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    asistencia = Asistencia.objects.filter(estudiante=estudiante).order_by('fecha')
    return render(request, 'reporte_asistencia_estudiante.html', {'estudiante': estudiante, 'asistencia': asistencia})

def generar_word_asistencia_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    asistencia = Asistencia.objects.filter(estudiante=estudiante).order_by('fecha')

    # Crear el documento de Word
    document = Document()
    document.add_heading(f'Reporte de Asistencia de {estudiante.nombre}', 0)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asignatura'
    hdr_cells[1].text = 'Fecha'
    hdr_cells[2].text = 'Presente'

    for registro in asistencia:
        row_cells = table.add_row().cells
        row_cells[0].text = registro.asignatura.nombre
        row_cells[1].text = str(registro.fecha)
        row_cells[2].text = 'Sí' if registro.presente else 'No'

    # Crear la carpeta si no existe
    output_dir = os.path.join(settings.MEDIA_ROOT, 'reportes_asistencia')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Guardar el documento final en una ruta temporal
    output_path = os.path.join(output_dir, f'reporte_asistencia_{estudiante.nombre}.docx')
    document.save(output_path)

    # Descargar el archivo Word generado
    with open(output_path, 'rb') as doc_file:
        response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="reporte_asistencia_{estudiante.nombre}.docx"'
        return response

def generar_word_notas_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    notas = Nota.objects.filter(estudiante=estudiante).order_by('asignatura__nombre')

    document = Document()
    document.add_heading(f'Reporte de Notas de {estudiante.nombre}', 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asignatura'
    hdr_cells[1].text = 'Nota'

    for nota in notas:
        row_cells = table.add_row().cells
        row_cells[0].text = nota.asignatura.nombre
        row_cells[1].text = str(nota.nota)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="reporte_notas_{estudiante.nombre}.docx"'
    document.save(response)
    return response


def generar_word_anotaciones_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    anotaciones = Anotacion.objects.filter(estudiante=estudiante).order_by('fecha')

    document = Document()
    document.add_heading(f'Reporte de Anotaciones de {estudiante.nombre}', 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fecha'
    hdr_cells[1].text = 'Anotación'

    for anotacion in anotaciones:
        row_cells = table.add_row().cells
        row_cells[0].text = str(anotacion.fecha)
        row_cells[1].text = anotacion.descripcion

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="reporte_anotaciones_{estudiante.nombre}.docx"'
    document.save(response)
    return response

def generar_word_certificado_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)

    # Ruta de la plantilla de documento de Word
    plantilla_path = os.path.join(settings.MEDIA_ROOT, 'plantilla_certificado_alumno.docx')

    # Intentar abrir el documento .docx
    try:
        doc = Document(plantilla_path)
    except Exception as e:
        return HttpResponse(f"Error al abrir el documento: {e}", status=500)

    # Reemplazar los campos en el documento
    for paragraph in doc.paragraphs:
        if '@nombre@' in paragraph.text:
            paragraph.text = paragraph.text.replace('@nombre@', estudiante.nombre)
        if '@rut@' in paragraph.text:
            paragraph.text = paragraph.text.replace('@rut@', estudiante.rut)

    # Crear la carpeta si no existe
    output_dir = os.path.join(settings.MEDIA_ROOT, 'certificados_alumnos')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Guardar el documento final en una ruta temporal
    output_path = os.path.join(output_dir, f'certificado_alumno_{estudiante.nombre}.docx')
    doc.save(output_path)

    # Descargar el archivo Word generado
    with open(output_path, 'rb') as doc_file:
        response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="certificado_alumno_{estudiante.nombre}.docx"'
        return response


def select_estudiante_reporte_notas_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_reporte_notas.html', {'estudiantes': estudiantes})


def reporte_notas_estudiante_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    notas = Nota.objects.filter(estudiante=estudiante).order_by('asignatura__nombre')
    return render(request, 'reporte_notas_estudiante.html', {'estudiante': estudiante, 'notas': notas})

def select_estudiante_reporte_asistencia_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_reporte_asistencia.html', {'estudiantes': estudiantes})

def reporte_anotaciones_estudiante_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    anotaciones = Anotacion.objects.filter(estudiante=estudiante).order_by('fecha')
    return render(request, 'reporte_anotaciones_estudiante.html', {'estudiante': estudiante, 'anotaciones': anotaciones})

def select_estudiante_certificado_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_certificado.html', {'estudiantes': estudiantes})

def select_estudiante_reporte_anotaciones_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_reporte_anotaciones.html', {'estudiantes': estudiantes})

def generar_word_reporte_completo_view(request, estudiante_id):
    estudiante = get_object_or_404(Estudiante, id=estudiante_id)
    notas = Nota.objects.filter(estudiante=estudiante).order_by('asignatura__nombre')
    asistencia = Asistencia.objects.filter(estudiante=estudiante).order_by('fecha')
    anotaciones = Anotacion.objects.filter(estudiante=estudiante).order_by('fecha')

    document = Document()
    document.add_heading(f'Reporte Completo de {estudiante.nombre}', 0)

    # Información del Estudiante
    document.add_heading('Información del Estudiante', level=1)
    document.add_paragraph(f'Nombre: {estudiante.nombre}')
    document.add_paragraph(f'Correo: {estudiante.correo}')
    document.add_paragraph(f'RUT: {estudiante.rut}')
    document.add_paragraph(f'Tipo: {estudiante.tipo}')

    # Notas
    document.add_heading('Notas', level=1)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asignatura'
    hdr_cells[1].text = 'Nota'
    for nota in notas:
        row_cells = table.add_row().cells
        row_cells[0].text = nota.asignatura.nombre
        row_cells[1].text = str(nota.nota)

    # Asistencia
    document.add_heading('Asistencia', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Asignatura'
    hdr_cells[1].text = 'Fecha'
    hdr_cells[2].text = 'Presente'
    for registro in asistencia:
        row_cells = table.add_row().cells
        row_cells[0].text = registro.asignatura.nombre
        row_cells[1].text = str(registro.fecha)
        row_cells[2].text = 'Sí' if registro.presente else 'No'

    # Anotaciones
    document.add_heading('Anotaciones', level=1)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fecha'
    hdr_cells[1].text = 'Descripción'
    for anotacion in anotaciones:
        row_cells = table.add_row().cells
        row_cells[0].text = str(anotacion.fecha)
        row_cells[1].text = anotacion.descripcion

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="reporte_completo_{estudiante.nombre}.docx"'
    document.save(response)
    return response

def select_estudiante_reporte_completo_view(request):
    estudiantes = Estudiante.objects.all()
    return render(request, 'select_estudiante_reporte_completo.html', {'estudiantes': estudiantes})

@login_required
def ver_asignaturas_notas(request):
    estudiante = request.user
    asignaturas = estudiante.asignaturas.all()
    notas = Nota.objects.filter(estudiante=estudiante)
    return render(request, 'ver_asignaturas_notas.html', {
        'asignaturas': asignaturas,
        'notas': notas,
    })

@login_required
def ver_anotaciones(request):
    estudiante = request.user
    anotaciones = Anotacion.objects.filter(estudiante=estudiante)
    return render(request, 'ver_anotaciones.html', {'anotaciones': anotaciones})

@login_required
def ver_eventos(request):
    eventos = Evento.objects.all()
    return render(request, 'ver_eventos.html', {'eventos': eventos})

def list_nodocentes_view(request):
    nodocentes = Nodocente.objects.all()
    return render(request, 'list_nodocentes.html', {'nodocentes': nodocentes})